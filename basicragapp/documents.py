from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE = 10 * 1024 * 1024


class DocumentError(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentError("Only PDF, DOCX, and TXT files are supported")
    if not content:
        raise DocumentError("The document is empty")
    if len(content) > MAX_FILE_SIZE:
        raise DocumentError("The document exceeds the 10 MB limit")

    try:
        if extension == ".pdf":
            text = "\n\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
        elif extension == ".docx":
            document = Document(BytesIO(content))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            blocks.extend(" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
            text = "\n\n".join(blocks)
        else:
            text = content.decode("utf-8-sig")
    except (OSError, UnicodeError, ValueError) as error:
        raise DocumentError(f"Could not read this {extension[1:].upper()} document") from error

    text = "\n".join(line.rstrip() for line in text.replace("\x00", "").splitlines()).strip()
    if not text:
        raise DocumentError("No readable text was found in the document")
    return text
