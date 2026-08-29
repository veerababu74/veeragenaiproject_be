import asyncio
import os
import unittest
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document
from fastapi import HTTPException

from basicragapp import bucket_storage, vector_store
from basicragapp.chunking import fixed_chunks, recursive_chunks, semantic_chunks
from basicragapp.documents import DocumentError, extract_text
from basicragapp.database import DuplicateDocumentError, RETENTION_SECONDS, StorageQuotaError, add_exchange, create_session, delete_session, document_storage_used, find_document_by_hash, get_session, initialize, list_documents, save_document
from basicragapp.embeddings import EMBEDDING_DIMENSION, EmbeddingError, embed_texts
from basicragapp.models import RagChatRequest
from basicragapp.router import _http_error, _rollback_document, send_message


class BasicRagFoundationTests(unittest.TestCase):
    def test_vercel_uses_writable_hugging_face_cache(self):
        with tempfile.TemporaryDirectory() as temporary_directory, patch.dict(
            os.environ, {"VERCEL": "1"}
        ), patch("basicragapp.bucket_storage.gettempdir", return_value=temporary_directory):
            bucket_storage._configure_cache()
            expected = Path(temporary_directory) / "veeragenai" / "huggingface"
            self.assertEqual(os.environ["HF_HOME"], str(expected))
            self.assertEqual(os.environ["HF_XET_CACHE"], str(expected / "xet"))

    @patch("basicragapp.embeddings.requests.post")
    def test_embedding_request_sets_pinecone_dimension(self, post):
        post.return_value.ok = True
        post.return_value.json.return_value = {"embeddings": [{"values": [0.0] * EMBEDDING_DIMENSION}]}

        vectors = embed_texts("test-key", "gemini-embedding-001", ["hello"], "RETRIEVAL_DOCUMENT")

        request = post.call_args.kwargs["json"]["requests"][0]
        self.assertEqual(request["taskType"], "RETRIEVAL_DOCUMENT")
        self.assertEqual(request["outputDimensionality"], EMBEDDING_DIMENSION)
        self.assertEqual(len(vectors[0]), EMBEDDING_DIMENSION)

    @patch("basicragapp.embeddings.requests.post")
    def test_embedding_dimension_mismatch_fails_before_storage(self, post):
        post.return_value.ok = True
        post.return_value.json.return_value = {"embeddings": [{"values": [0.0] * 3072}]}

        with self.assertRaisesRegex(EmbeddingError, "768-dimension"):
            embed_texts("test-key", "gemini-embedding-001", ["hello"], "RETRIEVAL_DOCUMENT")

    def test_accepts_txt_and_docx_but_rejects_other_types(self):
        self.assertEqual(extract_text("notes.txt", b"hello\r\nworld"), "hello\nworld")
        document = Document()
        document.add_paragraph("Project notes")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Value"
        stream = BytesIO()
        document.save(stream)
        self.assertIn("Name | Value", extract_text("notes.docx", stream.getvalue()))
        with self.assertRaises(DocumentError):
            extract_text("notes.csv", b"unsafe")

    def test_fixed_chunks_show_requested_overlap(self):
        text = "0123456789" * 30
        chunks = fixed_chunks(text, 100, 20)
        self.assertEqual(chunks[0][-20:], chunks[1][:20])

    def test_recursive_chunks_respect_boundaries_and_overlap(self):
        text = "First paragraph. " * 10 + "\n\n" + "Second paragraph. " * 10
        chunks = recursive_chunks(text, 120, 15)
        self.assertGreater(len(chunks), 2)
        self.assertTrue(all(len(chunk) <= 135 for chunk in chunks))
        self.assertEqual(chunks[0][-15:].strip(), chunks[1][:15].strip())

    def test_semantic_chunks_split_when_topic_similarity_drops(self):
        chunks = semantic_chunks(
            ["Cats purr.", "Kittens purr.", "Engines need oil."],
            [[1.0, 0.0], [0.9, 0.1], [0.0, 1.0]],
        )
        self.assertEqual(chunks, ["Cats purr. Kittens purr.", "Engines need oil."])


class BasicRagOwnershipTests(unittest.TestCase):
    def setUp(self):
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.database_path = Path(self.temporary_directory.name) / "rag.db"
        initialize(self.database_path)

    def tearDown(self):
        self.temporary_directory.cleanup()

    def test_sessions_and_documents_are_scoped_to_the_owning_user(self):
        session = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        save_document("doc-1", session["id"], "notes.txt", "text/plain", 5, "fixed", 100, 10, "private/path", ["hello"], self.database_path)

        self.assertIsNotNone(get_session(session["id"], "user-1", self.database_path))
        self.assertIsNone(get_session(session["id"], "user-2", self.database_path))
        self.assertEqual(len(list_documents(session["id"], "user-1", True, self.database_path)), 1)
        self.assertEqual(list_documents(session["id"], "user-2", True, self.database_path), [])

    def test_exchange_returns_false_if_session_was_deleted(self):
        session = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        delete_session(session["id"], "user-1", self.database_path)

        self.assertFalse(add_exchange(session["id"], "question", "answer", [], self.database_path))

    def test_user_has_five_mb_document_allowance(self):
        session = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        one_mb = 1024 * 1024
        save_document("doc-1", session["id"], "one.txt", "text/plain", one_mb, "fixed", 100, 10, "private/one", ["one"], self.database_path)

        self.assertEqual(document_storage_used("user-1", self.database_path), one_mb)
        save_document("doc-2", session["id"], "four.txt", "text/plain", 4 * one_mb, "fixed", 100, 10, "private/four", ["four"], self.database_path)
        with self.assertRaisesRegex(StorageQuotaError, "0.00 MB"):
            save_document("doc-3", session["id"], "extra.txt", "text/plain", 1, "fixed", 100, 10, "private/extra", ["extra"], self.database_path)

    def test_rag_sessions_expire_24_hours_after_last_update(self):
        with patch("basicragapp.database.time.time", return_value=100):
            session = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        with patch("basicragapp.database.time.time", return_value=1000):
            add_exchange(session["id"], "question", "answer", [], self.database_path)
        with patch("basicragapp.database.time.time", return_value=1000 + RETENTION_SECONDS - 1):
            self.assertIsNotNone(get_session(session["id"], "user-1", self.database_path))
        with patch("basicragapp.database.time.time", return_value=1000 + RETENTION_SECONDS):
            self.assertIsNone(get_session(session["id"], "user-1", self.database_path))

    def test_document_hash_is_scoped_to_user_across_sessions(self):
        first = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        second = create_session("user-1", "gemini", "gemini-test", "gemini-embedding-001", self.database_path)
        save_document("doc-1", first["id"], "notes.txt", "text/plain", 5, "fixed", 100, 10, "private/path", ["hello"], self.database_path, content_hash="abc123")

        duplicate = find_document_by_hash("user-1", "abc123", self.database_path)

        self.assertEqual(duplicate["filename"], "notes.txt")
        self.assertIsNone(find_document_by_hash("user-2", "abc123", self.database_path))
        with self.assertRaises(DuplicateDocumentError):
            save_document("doc-2", second["id"], "copy.txt", "text/plain", 5, "fixed", 100, 10, "private/copy", ["hello"], self.database_path, content_hash="abc123")


class BasicRagRemoteStoreTests(unittest.TestCase):
    @patch("basicragapp.router.database.add_exchange", return_value=False)
    @patch("basicragapp.router.database.get_messages", return_value=[])
    @patch("basicragapp.router.database.list_documents", return_value=[{"id": "doc-1"}])
    @patch("basicragapp.router.database.get_session")
    @patch("basicragapp.router.chat", return_value="answer")
    @patch("basicragapp.router.vector_store.query", return_value=[])
    @patch("basicragapp.router.embed_texts", return_value=[[0.0] * EMBEDDING_DIMENSION])
    def test_message_returns_conflict_if_session_is_deleted_during_generation(
        self, _embed, _query, _chat, get_session_mock, _documents, _messages, _add_exchange,
    ):
        get_session_mock.return_value = {
            "id": "session-1", "provider": "gemini", "model": "gemini-test",
            "embedding_model": "gemini-embedding-001",
        }
        request = RagChatRequest(
            session_id="session-1", provider="gemini", api_key="test-key",
            model="gemini-test", embedding_api_key="embed-key",
            embedding_model="gemini-embedding-001", message="question",
        )

        with self.assertRaises(HTTPException) as raised:
            asyncio.run(send_message(request, "user-1"))

        self.assertEqual(raised.exception.status_code, 409)

    def test_remote_failure_is_logged(self):
        try:
            raise vector_store.VectorStoreError("index unavailable")
        except vector_store.VectorStoreError as error:
            with self.assertLogs("veera.rag", level="ERROR") as logs:
                response = _http_error(error)

        self.assertEqual(response.status_code, 502)
        self.assertIn("RAG operation failed", logs.output[0])

    @patch("basicragapp.router.bucket_storage.delete")
    @patch("basicragapp.router.vector_store.delete_document")
    def test_failed_ingestion_attempts_both_remote_rollbacks(self, delete_vectors, delete_original):
        delete_vectors.side_effect = vector_store.VectorStoreError("failed")

        _rollback_document("user-1", "doc-1", 2, "private/doc-1.txt")

        delete_vectors.assert_called_once_with("user-1", "doc-1", 2)
        delete_original.assert_called_once_with("private/doc-1.txt")

    @patch("basicragapp.vector_store._index")
    def test_query_uses_user_namespace_and_session_filter(self, index_factory):
        index = Mock()
        index.query.return_value.matches = [Mock(score=0.9, metadata={"text": "hello"})]
        index_factory.return_value = index

        self.assertEqual(vector_store.query("user-1", "session-1", [0.1], 5)[0]["text"], "hello")
        index.query.assert_called_once_with(
            namespace="user-1", vector=[0.1], top_k=5, include_values=False,
            include_metadata=True, filter={"session_id": {"$eq": "session-1"}},
        )

    @patch("basicragapp.vector_store._index")
    def test_delete_uses_document_vector_ids_in_user_namespace(self, index_factory):
        index = Mock()
        index_factory.return_value = index

        vector_store.delete_document("user-1", "doc-1", 3)

        index.delete.assert_called_once_with(
            ids=["doc-1:0", "doc-1:1", "doc-1:2"], namespace="user-1",
        )


if __name__ == "__main__":
    unittest.main()
